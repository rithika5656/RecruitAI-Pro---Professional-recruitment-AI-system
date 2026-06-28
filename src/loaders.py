"""File loaders for JSON, JSONL, DOCX, and CSV inputs."""

from __future__ import annotations

import csv
import json
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

from docx import Document

from .utils import load_json, normalize_whitespace


@dataclass(slots=True)
class LoadedDocument:
    """Container for loaded source content."""

    path: Path
    content: Any
    file_type: str


class DataLoader:
    """Load source files into Python structures for downstream processing."""

    def load(self, path: Path) -> LoadedDocument:
        suffix = path.suffix.lower()
        if suffix == ".jsonl":
            return LoadedDocument(path=path, content=self._load_jsonl(path), file_type="jsonl")
        if suffix == ".json":
            return LoadedDocument(path=path, content=load_json(path), file_type="json")
        if suffix == ".csv":
            return LoadedDocument(path=path, content=self._load_csv(path), file_type="csv")
        if suffix == ".docx":
            return LoadedDocument(path=path, content=self._load_docx(path), file_type="docx")
        raise ValueError(f"Unsupported file type: {path.suffix}")

    def _load_jsonl(self, path: Path) -> list[dict[str, Any]]:
        """Load a JSONL file line by line."""

        records: list[dict[str, Any]] = []
        for raw_line in path.read_text(encoding="utf-8", errors="ignore").splitlines():
            line = raw_line.strip()
            if not line:
                continue
            records.append(json.loads(line))
        return records

    def _load_csv(self, path: Path) -> list[dict[str, str]]:
        """Load a CSV file as a list of dictionaries."""

        with path.open("r", encoding="utf-8", errors="ignore", newline="") as handle:
            reader = csv.DictReader(handle)
            return [dict(row) for row in reader]

    def _load_docx(self, path: Path) -> str:
        """Extract paragraph text from a DOCX file."""

        document = Document(str(path))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        return normalize_whitespace("\n".join(paragraphs))

    def load_many(self, paths: Iterable[Path]) -> list[LoadedDocument]:
        """Load multiple files in one call."""

        return [self.load(path) for path in paths]
