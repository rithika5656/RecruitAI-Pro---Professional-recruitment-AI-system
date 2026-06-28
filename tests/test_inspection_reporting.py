from __future__ import annotations

from pathlib import Path

from docx import Document

from src.inspection_cli import build_report, discover_input_files, inspect_input_directory
from src.inspection_reporting import InspectionReporter


def test_inspection_cli_discovers_and_reports_supported_files(tmp_path: Path) -> None:
    jsonl_path = tmp_path / "candidates.jsonl"
    jsonl_path.write_text('{"candidate_id": "c1", "skills": ["Python", "SQL"]}\n', encoding="utf-8")

    json_path = tmp_path / "candidate_schema.json"
    json_path.write_text('{"candidate_id": "string", "skills": "list"}', encoding="utf-8")

    doc = Document()
    doc.add_paragraph("Senior Data Scientist")
    doc.add_paragraph("Location: Remote")
    docx_path = tmp_path / "job_description.docx"
    doc.save(docx_path)

    discovered = discover_input_files(tmp_path)
    assert discovered == [json_path, jsonl_path, docx_path]

    reports = inspect_input_directory(tmp_path)
    reporter = InspectionReporter()
    markdown = reporter.build_markdown_report(reports)

    assert len(reports) == 3
    assert "candidates.jsonl" in markdown
    assert "candidate_schema.json" in markdown
    assert "job_description.docx" in markdown


def test_inspection_cli_writes_markdown_output(tmp_path: Path) -> None:
    candidates_path = tmp_path / "candidates.jsonl"
    candidates_path.write_text('{"candidate_id": "c1", "skills": ["Python"]}\n', encoding="utf-8")

    output_path = tmp_path / "inspection_report.md"
    result = build_report(tmp_path, output_path)

    assert result.output_path == output_path
    assert output_path.exists()
    assert output_path.read_text(encoding="utf-8").startswith("# Dataset Inspection Report")